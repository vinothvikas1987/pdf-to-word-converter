import os
import sys
from pdf2image import convert_from_path
import easyocr
from docx import Document
from tkinter import Tk, filedialog, messagebox

def main():
    Tk().withdraw()

    pdf_path = filedialog.askopenfilename(title="Select PDF File", filetypes=[("PDF Files", "*.pdf")])
    if not pdf_path:
        messagebox.showerror("Error", "No PDF file selected.")
        return

    output_docx = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if not output_docx:
        messagebox.showerror("Error", "No output file specified.")
        return

    # poppler_path = os.path.join(os.getcwd(), "poppler", "bin")
    def resource_path(relative_path):
    # """ Get absolute path to resource, works for dev and PyInstaller """
        try:
            base_path = sys._MEIPASS  # when running from PyInstaller bundle
        except AttributeError:
            base_path = os.getcwd()   # when running from source
        return os.path.join(base_path, relative_path)

    poppler_path = resource_path("poppler/bin")


    if not os.path.exists(os.path.join(poppler_path, "pdftoppm.exe")):
        messagebox.showerror("Error", f"Poppler not found at {poppler_path}")
        return

    reader = easyocr.Reader(['en'], gpu=False)

    try:
        pages = convert_from_path(pdf_path, dpi=400, poppler_path=poppler_path)
        if not pages:
            messagebox.showerror("Error", "No pages extracted from the PDF.")
            return
    except Exception as e:
        messagebox.showerror("Error", f"PDF conversion failed:\n{e}")
        return

    doc = Document()

    for i, page in enumerate(pages):
        image_path = f"temp_page_{i + 1}.jpg"
        try:
            page.save(image_path, 'JPEG')
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save image for page {i+1}:\n{e}")
            continue

        results = reader.readtext(image_path, detail=0, paragraph=True)
        text = "\n".join(results)

        doc.add_paragraph(f'--- Page {i + 1} ---')
        doc.add_paragraph(text)
        doc.add_paragraph('\n')

        if os.path.exists(image_path):
            os.remove(image_path)

    doc.save(output_docx)
    messagebox.showinfo("Success", f"OCR complete!\nSaved to: {output_docx}")

if __name__ == "__main__":
    main()


# import os
# from pdf2image import convert_from_path
# import easyocr
# from docx import Document
# from tkinter import filedialog, Tk


# # === Setup ===
# # pdf_path = r"D:\OCR\6th.pdf"
# # output_docx = r"D:\OCR\output_easyocr.docx"
# # # poppler_path = r"D:\poppler\poppler-24.08.0\Library\bin"


# # Hide root window
# Tk().withdraw()

# pdf_path = filedialog.askopenfilename(title="Select PDF File", filetypes=[("PDF Files", "*.pdf")])
# if not pdf_path:
#     print("❌ No file selected.")
#     exit()

# output_docx = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
# if not output_docx:
#     print("❌ No output file selected.")
#     exit()

#     poppler_path = os.path.join(os.getcwd(), "poppler", "bin")


# # === Init EasyOCR reader (only English) ===
# reader = easyocr.Reader(['en'])

# # === Convert PDF pages to images ===
# print("Converting PDF to images...")
# pages = convert_from_path(pdf_path, dpi=400, poppler_path=poppler_path)

# # === Create a new Word doc ===
# doc = Document()

# # === OCR each page ===
# for i, page in enumerate(pages):
#     print(f"Processing page {i + 1}...")
#     image_path = f"temp_page_{i + 1}.jpg"
#     page.save(image_path, 'JPEG')

#     results = reader.readtext(image_path, detail=0, paragraph=True)
#     text = "\n".join(results)

#     doc.add_paragraph(f'--- Page {i + 1} ---')
#     doc.add_paragraph(text)
#     doc.add_paragraph('\n')

#     os.remove(image_path)  # clean up temp file

# # === Save Word doc ===
# doc.save(output_docx)
# # print(f"\n✅ EasyOCR complete. Saved output to: {output_docx}")

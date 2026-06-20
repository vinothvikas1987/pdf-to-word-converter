import os
import sys
import threading
import easyocr
from tkinter import Tk, Label, Button, filedialog, StringVar, messagebox
from tkinter import ttk
from pdf2image import convert_from_path
from docx import Document

def resource_path(relative_path):
    """Support for PyInstaller frozen builds"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.getcwd()
    return os.path.join(base_path, relative_path)

class OCRApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Convert to Word from PDF")
        self.root.geometry("520x320")
        self.root.resizable(False, False)

        # State variables
        self.pdf_path = None
        self.output_path = None
        self.reader = None
        self.poppler_path = resource_path("poppler/bin")

        self.status_text = StringVar()
        self.status_text.set("📄 Please upload a PDF to begin.")

        # UI Layout
        Label(root, text="PDF to Word using EasyOCR", font=("Helvetica", 16)).pack(pady=10)
        Label(root, textvariable=self.status_text, wraplength=460, justify="center").pack(pady=10)

        self.progress = ttk.Progressbar(root, orient="horizontal", mode="determinate", length=450)
        self.progress.pack(pady=5)

        Button(root, text="📁 Upload PDF", command=self.select_pdf, width=25).pack(pady=5)
        Button(root, text="💾 Select Save Location", command=self.select_output_path, width=25).pack(pady=5)
        Button(root, text="▶️ Start OCR", command=self.start_ocr, width=25).pack(pady=5)
        Button(root, text="❌ Exit", command=root.quit, width=25).pack(pady=10)

    def select_pdf(self):
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if path:
            self.pdf_path = path
            self.status_text.set(f"✅ Selected PDF: {os.path.basename(path)}")

    def select_output_path(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile="output_easyocr.docx"
        )
        if path:
            self.output_path = path
            self.status_text.set("📍 Output path set.")

    def start_ocr(self):
        if not self.pdf_path:
            messagebox.showerror("Error", "Please select a PDF file.")
            return
        if not self.output_path:
            messagebox.showerror("Error", "Please set the output DOCX file path.")
            return
        if not os.path.exists(os.path.join(self.poppler_path, "pdftoppm.exe")):
            messagebox.showerror("Error", f"Poppler not found at:\n{self.poppler_path}")
            return

        # Run OCR in a separate thread
        threading.Thread(target=self.run_ocr).start()

    def run_ocr(self):
        try:
            self.status_text.set("🧠 Loading EasyOCR model...")
            self.reader = easyocr.Reader(['en'], gpu=False)

            self.status_text.set("🖼️ Converting PDF to images...")
            pages = convert_from_path(self.pdf_path, dpi=400, poppler_path=self.poppler_path)
            total_pages = len(pages)
            self.progress["maximum"] = total_pages

            doc = Document()

            for i, page in enumerate(pages):
                self.status_text.set(f"🔍 OCR: Page {i + 1} of {total_pages}")
                self.progress["value"] = i + 1
                self.root.update_idletasks()

                image_path = f"temp_page_{i + 1}.jpg"
                page.save(image_path, 'JPEG')

                results = self.reader.readtext(image_path, detail=0, paragraph=True)
                text = "\n".join(results)

                doc.add_paragraph(f'--- Page {i + 1} ---')
                doc.add_paragraph(text)
                doc.add_paragraph('\n')

                os.remove(image_path)

            doc.save(self.output_path)
            self.status_text.set(f"✅ completed! Saved to:\n{self.output_path}")
            messagebox.showinfo("Success", f"Conversion completed!\nSaved to:\n{self.output_path}")

        except Exception as e:
            self.status_text.set("❌ An error occurred.")
            messagebox.showerror("OCR Error", str(e))

def main():
    root = Tk()
    app = OCRApp(root)
    root.after(100, lambda: None)  # Force UI to render first
    root.mainloop()

if __name__ == "__main__":
    main()
